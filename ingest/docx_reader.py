import docx
from docx.oxml.ns import qn

# Word stores text-box content twice; this tag marks the duplicate legacy copy.
MC_FALLBACK = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"


def _is_in_fallback(node) -> bool:
    """True if this text node is inside a legacy VML fallback (a duplicate)."""
    ancestor = node.getparent()
    while ancestor is not None:
        if ancestor.tag == MC_FALLBACK:
            return True
        ancestor = ancestor.getparent()
    return False


def extract_text_from_docx(path: str) -> str:
    document = docx.Document(path)
    parts = []

    # 1. Normal paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # 2. Tables (CV templates often store content here)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    text = "\n".join(parts)

    # 3. Fallback for text boxes / shapes — skip the duplicate legacy copies
    if not text.strip():
        nodes = document.element.body.iter(qn("w:t"))
        parts = [n.text for n in nodes if n.text and not _is_in_fallback(n)]
        text = "\n".join(parts)

    return text